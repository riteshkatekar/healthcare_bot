#service.py

from __future__ import annotations

import base64
import csv
import json
import os
import re
import sqlite3
import tempfile
import time
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from groq import Groq
from PIL import Image, ImageOps
from pypdf import PdfReader

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from langdetect import detect as langdetect_detect
except Exception:
    langdetect_detect = None


ALLOWED_TEXT_EXTS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".html", ".htm", ".xml", ".log"}
ALLOWED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}
ALLOWED_AUDIO_EXTS = {".wav", ".mp3", ".m4a", ".mp4", ".mpeg", ".mpga", ".ogg", ".webm", ".flac"}

WHITESPACE_RE = re.compile(r"\s+")
HTML_TAG_RE = re.compile(r"<[^>]+>")
DEVANAGARI_RE = re.compile(r"[\u0900-\u097F]")
BENGALI_RE = re.compile(r"[\u0980-\u09FF]")
GURMUKHI_RE = re.compile(r"[\u0A00-\u0A7F]")
GUJARATI_RE = re.compile(r"[\u0A80-\u0AFF]")
TAMIL_RE = re.compile(r"[\u0B80-\u0BFF]")
TELUGU_RE = re.compile(r"[\u0C00-\u0C7F]")
KANNADA_RE = re.compile(r"[\u0C80-\u0CFF]")
MALAYALAM_RE = re.compile(r"[\u0D00-\u0D7F]")

JSON_BLOCK_RE = re.compile(r"\{.*\}", re.DOTALL)
LATIN_RE = re.compile(r"[A-Za-z]")
MIXED_TOKEN_RE = re.compile(r"(?=.*[A-Za-z])(?=.*[\u0900-\u0DFF])")
TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9\-./]*|[\u0900-\u0DFF][\u0900-\u0DFF\-./]*")

HI_HINTS = {
    "hai", "kya", "kyu", "kyun", "dard", "bukhar", "khansi", "sardi",
    "pet", "sar", "doctor", "dawai", "kripya", "please", "nahi", "ho",
}
MR_HINTS = {
    "आहे", "काय", "का", "दुखत", "ताप", "खोकला", "डॉक्टर", "औषध",
    "कृपया", "होत", "नाही", "मला", "सांग", "सांगू",
}

EMERGENCY_KEYWORDS = {
    "chest pain", "difficulty breathing", "shortness of breath",
    "seizure", "stroke", "unconscious", "fainting", "blue lips",
    "severe bleeding", "suicidal", "self harm", "overdose",
    "heart attack", "anaphylaxis", "confusion", "one-sided weakness",
    "sudden vision loss", "severe allergic reaction",
    "छाती दुख", "श्वास घेण्यास त्रास", "बेशुद्ध", "अचानक कमजोरी",
    "आत्महत्या", "स्वतःला इजा", "तीव्र रक्तस्राव",
    "सीने में दर्द", "सांस लेने में दिक्कत", "बेहोशी", "आत्महत्या",
}

COMMON_CONTEXT_WORDS = {
    "tablet", "tab", "capsule", "caps", "syrup", "dose", "dosage", "mg", "ml", "injection",
    "medicine", "medication", "drug", "antibiotic", "antibiotics", "painkiller", "analgesic",
    "fever", "cold", "cough", "pain", "headache", "stomach", "nausea", "vomiting", "diarrhea",
    "allergy", "allergic", "prescribed", "prescription", "take", "taken", "take it", "use",
    "should", "avoid", "side effect", "side effects", "doctor", "hospital", "clinic",
    "ताप", "औषध", "गोळी", "सिरप", "डोस", "इंजेक्शन", "वेदना", "डोकेदुखी",
    "fever", "bukhar", "khansi", "sardi", "दवाई", "डोके", "पोट",
}

MEDICATION_SUFFIXES = (
    "mab", "nib", "vir", "cillin", "mycin", "azole", "statin", "olol", "pril", "sartan",
    "prazole", "dine", "azine", "oxetine", "caine", "sone", "lone", "tidine", "fen",
    "acetamol", "olam", "pam", "lam", "peridone", "triptyline", "barbital", "cort", "pred",
)

EN_STOPWORDS = {
    "the", "and", "or", "but", "if", "then", "also", "with", "without", "for", "to", "in", "on",
    "at", "by", "from", "of", "a", "an", "is", "are", "was", "were", "be", "been", "it", "this",
    "that", "these", "those", "you", "your", "we", "they", "he", "she", "them", "his", "her",
    "my", "me", "our", "their", "as", "not", "no", "do", "does", "did", "can", "could", "will",
    "would", "should", "may", "might", "must", "have", "has", "had", "take", "taken", "use",
    "used", "about", "because", "when", "what", "which", "who", "whom", "where", "why", "how",
    "please", "thanks", "thank", "okay", "ok",
}


def env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except Exception:
        return default


def env_str(name: str, default: str) -> str:
    value = os.getenv(name, default)
    return (value or "").strip() or default


def clean_text(text: str) -> str:
    text = text or ""
    text = text.replace("\x00", " ")
    text = WHITESPACE_RE.sub(" ", text)
    return text.strip()


def safe_snippet(text: str, limit: int = 12000) -> str:
    text = clean_text(text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + " ...[truncated]"


def safe_json_loads(text: str) -> Dict[str, Any]:
    text = clean_text(text)
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text, flags=re.IGNORECASE)
    try:
        return json.loads(text)
    except Exception:
        pass

    match = JSON_BLOCK_RE.search(text)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            pass

    return {}


def detect_language(text: str) -> str:
    text = clean_text(text)
    if not text:
        return "en"

    lower = text.lower()
    hi_score = sum(lower.count(tok) for tok in HI_HINTS)
    mr_score = sum(text.count(tok) for tok in MR_HINTS)

    if langdetect_detect and len(text) >= 20:
        try:
            lang = langdetect_detect(text)
            if lang.startswith("hi"):
                return "hi"
            if lang.startswith("mr"):
                return "mr"
            if lang.startswith("en"):
                return "en"
        except Exception:
            pass

    has_dev = bool(DEVANAGARI_RE.search(text))
    has_latin = bool(LATIN_RE.search(text))

    if has_dev and has_latin:
        return "mixed"
    if has_dev:
        return "mr" if mr_score > hi_score else "hi"

    if hi_score > 0 and mr_score > 0:
        return "mixed"
    if mr_score > hi_score and mr_score > 0:
        return "mr"
    if hi_score > 0:
        return "hi"

    return "en"


def is_emergency_text(text: str) -> bool:
    text = clean_text(text).lower()
    if not text:
        return False
    return any(keyword in text for keyword in EMERGENCY_KEYWORDS)


def chunk_text(text: str, chunk_size: int = 7000, overlap: int = 500) -> List[str]:
    text = clean_text(text)
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    total = len(text)
    while start < total:
        end = min(start + chunk_size, total)
        chunks.append(text[start:end])
        if end >= total:
            break
        start = max(0, end - overlap)
    return chunks


def extract_text_from_file(file_path: str, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    try:
        if suffix == ".pdf":
            reader = PdfReader(file_path)
            pages: List[str] = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    pages.append("")
            return clean_text("\n".join(pages))

        if suffix == ".docx":
            doc = Document(file_path)
            parts: List[str] = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return clean_text("\n".join(parts))

        if suffix == ".txt":
            return clean_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"))

        if suffix == ".md":
            return clean_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"))

        if suffix == ".csv":
            rows: List[str] = []
            with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    row = [cell.strip() for cell in row if cell is not None and str(cell).strip()]
                    if row:
                        rows.append(" | ".join(row))
            return clean_text("\n".join(rows))

        if suffix == ".json":
            raw = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            try:
                parsed = json.loads(raw)
                return clean_text(json.dumps(parsed, indent=2, ensure_ascii=False))
            except Exception:
                return clean_text(raw)

        if suffix in {".html", ".htm", ".xml", ".log"}:
            raw = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            if suffix in {".html", ".htm"}:
                raw = HTML_TAG_RE.sub(" ", raw)
            return clean_text(raw)

    except Exception:
        return ""

    return ""


def resize_image_for_vision(input_path: str, max_side: int = 1600) -> str:
    img = Image.open(input_path)
    img = ImageOps.exif_transpose(img).convert("RGB")

    w, h = img.size
    longest = max(w, h)
    if longest > max_side:
        scale = max_side / float(longest)
        new_size = (max(1, int(w * scale)), max(1, int(h * scale)))
        img = img.resize(new_size, Image.LANCZOS)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    tmp_path = tmp.name
    tmp.close()
    img.save(tmp_path, format="JPEG", quality=92, optimize=True)
    return tmp_path


def image_to_data_url(image_path: str) -> str:
    suffix = Path(image_path).suffix.lower().lstrip(".")
    mime = "jpeg" if suffix in {"jpg", "jpeg"} else suffix or "jpeg"
    with open(image_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")
    return f"data:image/{mime};base64,{encoded}"


def local_ocr_text(image_path: str) -> str:
    if pytesseract is None:
        return ""
    try:
        img = Image.open(image_path)
        img = ImageOps.exif_transpose(img).convert("RGB")
        text = pytesseract.image_to_string(img)
        return clean_text(text)
    except Exception:
        return ""


def infer_response_style(user_text: str) -> str:
    text = clean_text(user_text).lower()

    short_patterns = [
        r"\bshort\b",
        r"\bbrief\b",
        r"\bconcise\b",
        r"\bin 3 lines\b",
        r"\b3 lines\b",
        r"\bthree lines\b",
        r"\b2 lines\b",
        r"\bone line\b",
        r"\bquick\b",
        r"\bvery short\b",
    ]
    long_patterns = [
        r"\bdetailed\b",
        r"\blong\b",
        r"\bvery long\b",
        r"\bexplain in detail\b",
        r"\bdeep explanation\b",
        r"\bmore detail\b",
        r"\bin depth\b",
    ]

    if any(re.search(p, text) for p in short_patterns):
        return "short"
    if any(re.search(p, text) for p in long_patterns):
        return "detailed"
    return "concise"


def _lang_to_script(lang: str) -> str:
    lang = (lang or "").lower()
    if lang.startswith(("hi", "mr")):
        return "devanagari"
    if lang.startswith("kn"):
        return "kannada"
    if lang.startswith("ta"):
        return "tamil"
    if lang.startswith("te"):
        return "telugu"
    if lang.startswith("ml"):
        return "malayalam"
    if lang.startswith("bn"):
        return "bengali"
    if lang.startswith("gu"):
        return "gujarati"
    if lang.startswith("pa"):
        return "gurmukhi"
    return "latin"


def build_system_prompt(language: str, style: str = "concise") -> str:
    if style == "short":
        length_rule = "Answer in 2–4 short lines OR 3–4 bullet points maximum."
    elif style == "detailed":
        length_rule = "Give clear and useful explanation using short bullet points or small paragraphs."
    else:
        length_rule = "Keep the answer short and practical. Prefer 3–6 bullet points."

    return f"""
You are a practical and helpful healthcare assistant.

Core behavior:
- Always give useful, actionable advice first.
- Focus on home remedies, basic care, and simple guidance.
- Be direct, clear, and to the point.
- Avoid unnecessary warnings or generic disclaimers.

Safety behavior:
- Do NOT say "I am not a doctor".
- Do NOT refuse simple medical questions.
- Only add a short safety note at the END if needed.
- Mention doctor/hospital ONLY if symptoms are serious, high-risk, or persistent.

Language behavior:
- Reply entirely in the user's language.
- If the user mixes languages, reply naturally in that same mix.
- Do not switch back to English unless the user used English.
- If the language is not English, avoid random English words in ordinary sentences.
- If a medicine or drug name must appear, render it naturally for the target language.
- On first mention of a medicine or drug, keep the local-language form followed by the English original in parentheses.
- Never mix scripts inside one word.
- If the user asks for a specific language, honor it strictly.

Response format:
- Use clean bullet points when listing advice.
- Each bullet must be on a new line.
- Do NOT mix bullets with paragraphs.
- Keep spacing clean and readable.

Context usage:
- If image/file context is provided, use it properly.
- If unclear, say it briefly and continue with best possible help.

Style rule:
- {length_rule}

Language hint: {language}
""".strip()


def build_user_payload(
    user_message: str,
    file_context: str = "",
    image_context: str = "",
    image_ocr: str = "",
) -> str:
    parts: List[str] = []

    if user_message:
        parts.append(f"User message:\n{user_message.strip()}")

    if file_context.strip():
        parts.append(
            "Uploaded file context:\n"
            + file_context.strip()
            + "\n\nUse the uploaded document as grounding evidence."
        )

    if image_context.strip():
        parts.append(
            "Uploaded image context:\n"
            + image_context.strip()
            + "\n\nUse the image context carefully and mention uncertainty if needed."
        )

    if image_ocr.strip():
        parts.append("OCR text from image:\n" + image_ocr.strip())

    parts.append(
        "Answer safely. If medical advice is uncertain, say so and suggest a clinician when appropriate."
    )
    return "\n\n".join(parts)


def build_chat_messages(
    user_message: str,
    language: str,
    recent_messages: Sequence[Dict[str, str]],
    memory_summary: str = "",
    file_context: str = "",
    image_context: str = "",
    image_ocr: str = "",
) -> List[Dict[str, Any]]:
    style = infer_response_style(user_message)
    messages: List[Dict[str, Any]] = [
        {"role": "system", "content": f"TARGET_LANGUAGE_CODE={language}"},
        {"role": "system", "content": build_system_prompt(language, style)},
    ]

    if memory_summary.strip():
        messages.append(
            {
                "role": "system",
                "content": f"Conversation memory summary:\n{memory_summary.strip()}",
            }
        )

    for msg in recent_messages:
        role = msg.get("role", "")
        content = clean_text(msg.get("content", ""))
        if role in {"user", "assistant"} and content:
            messages.append({"role": role, "content": content})

    messages.append(
        {
            "role": "user",
            "content": build_user_payload(
                user_message=user_message,
                file_context=file_context,
                image_context=image_context,
                image_ocr=image_ocr,
            ),
        }
    )
    return messages


@dataclass
class ImageInsight:
    filename: str = ""
    description: str = ""
    visible_text: str = ""
    medical_relevance: str = ""
    objects: List[str] = field(default_factory=list)
    notes: str = ""
    raw: Dict[str, Any] = field(default_factory=dict)

    def to_context_text(self) -> str:
        objects_text = ", ".join(self.objects) if self.objects else "None"
        return clean_text(
            f"Filename: {self.filename}\n"
            f"Description: {self.description or 'Not available'}\n"
            f"Objects: {objects_text}\n"
            f"Medical relevance: {self.medical_relevance or 'Not available'}\n"
            f"Visible text: {self.visible_text or 'None'}\n"
            f"Notes: {self.notes or 'None'}"
        )

    def to_short_response(self) -> str:
        parts: List[str] = []
        if self.description:
            parts.append(self.description)
        if self.visible_text:
            parts.append(f"Text seen: {self.visible_text}")
        if self.medical_relevance:
            parts.append(f"Medical note: {self.medical_relevance}")
        return clean_text(" ".join(parts)) or "Image processed successfully."


class MedicalTermNormalizer:
    def __init__(self, groq_service: "GroqService") -> None:
        self.groq = groq_service

    def _is_medical_candidate(self, token: str, context: str) -> bool:
        t = clean_text(token)
        if not t:
            return False

        low = t.lower().strip(".,;:!?()[]{}\"'`")
        if not low or low in EN_STOPWORDS:
            return False

        has_latin = bool(LATIN_RE.search(t))
        has_indic = bool(
            DEVANAGARI_RE.search(t)
            or BENGALI_RE.search(t)
            or GURMUKHI_RE.search(t)
            or GUJARATI_RE.search(t)
            or TAMIL_RE.search(t)
            or TELUGU_RE.search(t)
            or KANNADA_RE.search(t)
            or MALAYALAM_RE.search(t)
        )
        mixed = has_latin and has_indic

        if mixed:
            return True

        if any(low.endswith(suf) for suf in MEDICATION_SUFFIXES):
            return True

        if has_latin and len(low) >= 4:
            if any(cw in context.lower() for cw in COMMON_CONTEXT_WORDS):
                return True
            if low[0].isupper():
                return True
            if re.search(r"\d", low) or "-" in low:
                return True

        if len(low) >= 5 and (low in context.lower()):
            if any(cw in context.lower() for cw in COMMON_CONTEXT_WORDS):
                return True

        return False

    def extract_candidates(self, answer: str, source_text: str = "") -> List[str]:
        text = f"{answer or ''} {source_text or ''}".strip()
        if not text:
            return []

        candidates: List[str] = []
        seen: set[str] = set()

        for match in TOKEN_RE.finditer(text):
            token = match.group(0)
            if not self._is_medical_candidate(token, text):
                continue

            normalized = token.strip(".,;:!?()[]{}\"'`")
            if not normalized:
                continue

            key = normalized.lower()
            if key in seen:
                continue
            seen.add(key)
            candidates.append(normalized)

        for m in re.finditer(r"[\u0900-\u0DFF]+[A-Za-z][A-Za-z0-9\-./]*", text):
            token = m.group(0).strip(".,;:!?()[]{}\"'`")
            key = token.lower()
            if key not in seen:
                seen.add(key)
                candidates.append(token)

        return candidates[:20]

    def _needs_normalization(self, answer: str, target_language: str, candidates: List[str]) -> bool:
        if not answer or target_language.lower().startswith("en"):
            return False

        if candidates:
            return True

        if LATIN_RE.search(answer):
            return True

        return False

    def _rewrite_candidates_via_llm(
        self,
        answer: str,
        target_language: str,
        candidates: List[str],
        source_text: str = "",
    ) -> Dict[str, str]:
        if not self.groq.client:
            return {}

        prompt = f"""
You are a multilingual medical text normalizer.

Task:
- The response must be fully in {target_language}.
- Detect the following medical terms / drug names / mixed-script words.
- Return the best localized rendering for each term.
- For medicine/drug names, format as: Local-language-form (English original)
- Never return mixed-script words like "परacetamol".
- Keep meaning unchanged.
- Do not add new facts.
- Do not explain anything.

Return ONLY valid JSON in this exact format:
{{"terms": {{"original_term":"localized term (Original Term)"}}}}

Original response:
{answer}

Source context:
{source_text}

Terms to normalize:
{json.dumps(candidates, ensure_ascii=False)}
""".strip()

        try:
            raw = self.groq.chat(
                [
                    {"role": "system", "content": "You normalize medical text and return JSON only."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=280,
                retries=1,
                response_format={"type": "json_object"},
                postprocess=False,
            )
            data = safe_json_loads(raw)
            terms = data.get("terms", {})
            if isinstance(terms, dict):
                out: Dict[str, str] = {}
                for k, v in terms.items():
                    kk = clean_text(str(k))
                    vv = clean_text(str(v))
                    if kk and vv:
                        out[kk] = vv
                return out
        except Exception:
            return {}

        return {}

    def _apply_candidate_map(self, answer: str, mapping: Dict[str, str]) -> str:
        if not answer or not mapping:
            return answer

        out = answer

        for original in sorted(mapping.keys(), key=len, reverse=True):
            replacement = mapping[original]
            if not original or not replacement:
                continue

            pattern = re.compile(re.escape(original), re.IGNORECASE)

            def _repl(m: re.Match) -> str:
                matched = m.group(0)
                if replacement.lower() in matched.lower():
                    return matched
                return replacement

            out = pattern.sub(_repl, out)

        return out

    def _strong_rewrite(self, answer: str, target_language: str, source_text: str = "") -> str:
        if not self.groq.client:
            return answer

        prompt = f"""
Rewrite the text fully in {target_language}.

Hard rules:
- Keep the meaning exactly.
- Keep the same bullets/line breaks if present.
- Do not mix English into normal sentences.
- Any medicine or drug name must be written in the target language script followed by the English original in parentheses on first mention.
- Never produce mixed-script words like "परacetamol".
- Do not add or remove medical advice.
- Return plain text only.

Text:
{answer}

Context:
{source_text}
""".strip()

        try:
            return self.groq.chat(
                [
                    {"role": "system", "content": f"Rewrite only in {target_language} and keep all medical terms localized."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.15,
                max_tokens=450,
                retries=1,
                postprocess=False,
            )
        except Exception:
            return answer

    def normalize_answer(self, answer: str, target_language: str, source_text: str = "") -> str:
        answer = clean_text(answer)
        if not answer:
            return answer

        target_language = (target_language or "en").strip().lower()
        if target_language.startswith("en"):
            return answer

        candidates = self.extract_candidates(answer, source_text)
        if not self._needs_normalization(answer, target_language, candidates):
            return answer

        mapped = self._rewrite_candidates_via_llm(answer, target_language, candidates, source_text)
        if mapped:
            answer = self._apply_candidate_map(answer, mapped)

        if LATIN_RE.search(answer) or MIXED_TOKEN_RE.search(answer):
            answer = self._strong_rewrite(answer, target_language, source_text)

        answer = clean_text(answer)
        answer = re.sub(r"\s+([,.;:!?])", r"\1", answer)
        answer = re.sub(r"\(\s+", "(", answer)
        answer = re.sub(r"\s+\)", ")", answer)

        # Fix mixed-script words like परacetamol
        # 🔥 Step 1: Fix mixed-script words like परacetamol
        answer = re.sub(r"([\u0900-\u0DFF]+)([A-Za-z]+)", r"\1 \2", answer)

# 🔥 Step 2: Strong Marathi normalization (IMPORTANT)
        if target_language.startswith("mr") or target_language == "mixed":
         replacements = {
                "paracetamol": "पॅरासिटामोल (Paracetamol)",
                "crocin": "क्रोसिन (Crocin)",
                "dolo": "डोलो (Dolo)",
                "ibuprofen": "आयबुप्रोफेन (Ibuprofen)",
    }

         for eng, mar in replacements.items():
             pattern = re.compile(rf"\b{eng}\b", re.IGNORECASE)
             answer = pattern.sub(mar, answer)
        return answer


class GroqService:
    def __init__(self) -> None:
        self.api_key = env_str("GROQ_API_KEY", "")
        self.client = Groq(api_key=self.api_key) if self.api_key else None
        self.text_model = env_str("GROQ_TEXT_MODEL", "llama-3.3-70b-versatile")
        self.vision_model = env_str("GROQ_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")
        self.stt_model = env_str("GROQ_STT_MODEL", "whisper-large-v3-turbo")
        self.file_summary_threshold_chars = env_int("FILE_SUMMARY_THRESHOLD_CHARS", 12000)

        tesseract_cmd = env_str("TESSERACT_CMD", "")
        if tesseract_cmd and pytesseract is not None:
            try:
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            except Exception:
                pass

        self.normalizer = MedicalTermNormalizer(self)

    def _require_client(self) -> Groq:
        if self.client is None:
            raise RuntimeError("GROQ_API_KEY is missing. Please add it to your .env file.")
        return self.client

    def _extract_last_user_text(self, messages: List[Dict[str, Any]]) -> str:
        for msg in reversed(messages or []):
            if msg.get("role") == "user":
                content = msg.get("content", "")
                if isinstance(content, str):
                    return clean_text(content)
                if isinstance(content, list):
                    parts = []
                    for item in content:
                        if isinstance(item, dict) and item.get("type") == "text":
                            parts.append(str(item.get("text", "")))
                    return clean_text(" ".join(parts))
        return ""

    def _infer_style_from_messages(self, messages: List[Dict[str, Any]]) -> str:
        return infer_response_style(self._extract_last_user_text(messages))

    def _infer_target_language_from_messages(self, messages: List[Dict[str, Any]]) -> str:
        for msg in messages or []:
            if msg.get("role") != "system":
                continue
            content = clean_text(msg.get("content", ""))
            m = re.search(r"TARGET_LANGUAGE_CODE\s*=\s*([a-zA-Z\-]+)", content)
            if m:
                return m.group(1).strip().lower()
            m2 = re.search(r"Language hint:\s*([a-zA-Z\-]+)", content)
            if m2:
                return m2.group(1).strip().lower()
        return "en"

    def _token_cap(self, style: str) -> int:
        if style == "short":
            return 220
        if style == "detailed":
            return 650
        return 320

    def _chat_raw(
        self,
        messages: List[Dict[str, Any]],
        *,
        model: Optional[str] = None,
        temperature: float = 0.25,
        max_tokens: int = 900,
        response_format: Optional[Dict[str, Any]] = None,
        retries: int = 2,
    ) -> str:
        client = self._require_client()
        style = self._infer_style_from_messages(messages)
        bounded_tokens = min(max_tokens or 900, self._token_cap(style))

        kwargs: Dict[str, Any] = {
            "model": model or self.text_model,
            "messages": messages,
            "temperature": 0.2 if style != "detailed" else temperature,
            "top_p": 1,
            "max_completion_tokens": bounded_tokens,
        }
        if response_format:
            kwargs["response_format"] = response_format

        last_error: Optional[Exception] = None
        for attempt in range(retries + 1):
            try:
                resp = client.chat.completions.create(**kwargs)
                content = resp.choices[0].message.content if resp.choices else ""
                return clean_text(content or "")
            except Exception as exc:
                last_error = exc
                if attempt < retries:
                    time.sleep(0.7 * (attempt + 1))
                else:
                    break

        raise RuntimeError(clean_text(str(last_error)) or "Connection error.")

    def chat(
        self,
        messages: List[Dict[str, Any]],
        *,
        model: Optional[str] = None,
        temperature: float = 0.25,
        max_tokens: int = 900,
        response_format: Optional[Dict[str, Any]] = None,
        retries: int = 2,
        postprocess: bool = True,
        source_text: str = "",
    ) -> str:
        content = self._chat_raw(
            messages,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            response_format=response_format,
            retries=retries,
        )

        if not postprocess:
            return content

        target_language = self._infer_target_language_from_messages(messages)
        source = source_text or self._extract_last_user_text(messages)
        return self.normalizer.normalize_answer(content, target_language, source_text=source)

    def stream_chat(
        self,
        messages: List[Dict[str, Any]],
        *,
        model: Optional[str] = None,
        temperature: float = 0.25,
        max_tokens: int = 900,
        retries: int = 2,
    ):
        client = self._require_client()
        style = self._infer_style_from_messages(messages)
        bounded_tokens = min(max_tokens or 900, self._token_cap(style))

        kwargs: Dict[str, Any] = {
            "model": model or self.text_model,
            "messages": messages,
            "temperature": 0.2 if style != "detailed" else temperature,
            "top_p": 1,
            "max_completion_tokens": bounded_tokens,
            "stream": True,
        }

        last_error: Optional[Exception] = None
        for attempt in range(retries + 1):
            try:
                stream = client.chat.completions.create(**kwargs)
                for chunk in stream:
                    try:
                        delta = chunk.choices[0].delta.content or ""
                    except Exception:
                        delta = ""
                    if delta:
                        yield delta
                return
            except Exception as exc:
                last_error = exc
                if attempt < retries:
                    time.sleep(0.7 * (attempt + 1))
                else:
                    break

        raise RuntimeError(clean_text(str(last_error)) or "Connection error.")

    def summarize_text(self, text: str, language: str = "en", max_tokens: int = 350) -> str:
        text = clean_text(text)
        if not text:
            return ""

        prompt = f"""
Summarize this text for a healthcare chatbot.

Keep only:
- the main facts
- symptoms, dates, medications, numeric values, and risks
- any follow-up questions that should be asked

Write in the same language as the user's context: {language}

TEXT:
{text}
""".strip()

        return self.chat(
            [
                {"role": "system", "content": "You summarize content accurately and concisely."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.15,
            max_tokens=max_tokens,
            retries=1,
            postprocess=False,
        )

    def condense_document_context(self, text: str, language: str = "en") -> str:
        text = clean_text(text)
        if not text:
            return ""
        if len(text) <= self.file_summary_threshold_chars:
            return text
        return self.summarize_large_text(text, language=language)

    def summarize_large_text(self, text: str, language: str = "en") -> str:
        text = clean_text(text)
        if not text:
            return ""
        if len(text) <= self.file_summary_threshold_chars:
            return text

        chunks = chunk_text(text, chunk_size=7000, overlap=500)
        if not chunks:
            return ""

        partials: List[str] = []
        for chunk in chunks[:5]:
            try:
                partials.append(self.summarize_text(chunk, language=language, max_tokens=280))
            except Exception:
                partials.append(safe_snippet(chunk, 2200))

        merged = "\n\n".join(partials)
        if len(chunks) > 5:
            merged += "\n\n[More content was condensed because the file was large.]"

        try:
            final = self.summarize_text(merged, language=language, max_tokens=280)
            return final or merged
        except Exception:
            return merged

    def compress_chat_memory(self, previous_summary: str, older_messages: str, language: str = "en") -> str:
        previous_summary = clean_text(previous_summary)
        older_messages = clean_text(older_messages)

        prompt = f"""
You are compressing conversation memory for a healthcare chatbot.

Keep only useful long-term memory:
- symptoms and medical context
- user preferences
- ongoing questions
- important conclusions or warnings
- anything needed to continue the same conversation naturally

Do NOT add new medical advice.
Do NOT hallucinate facts.
Be concise.

Existing summary:
{previous_summary or "(none)"}

Older messages:
{older_messages}

Return a compact memory summary in the same language as the conversation context: {language}
""".strip()

        return self.chat(
            [
                {"role": "system", "content": "You compress conversation memory safely."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.15,
            max_tokens=280,
            retries=1,
            postprocess=False,
        )

    def generate_followups(
        self,
        *,
        user_message: str,
        assistant_answer: str,
        language: str = "en",
        memory_summary: str = "",
        file_context: str = "",
        image_context: str = "",
        image_ocr: str = "",
    ) -> List[str]:
        prompt = f"""
Generate 2 to 4 short follow-up questions for a healthcare chatbot.

Rules:
- Questions must be directly relevant to the last answer and user context.
- Keep them short, natural, and useful.
- Use the same language as the conversation: {language}
- Return ONLY valid JSON in this exact format:
  {{"questions":["question 1","question 2"]}}

Context:
User message: {user_message}
Assistant answer: {assistant_answer}
Memory summary: {memory_summary}
File context: {file_context}
Image context: {image_context}
OCR text: {image_ocr}
""".strip()

        raw = self.chat(
            [
                {"role": "system", "content": "You generate concise follow-up questions as JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=180,
            retries=1,
            response_format={"type": "json_object"},
            postprocess=False,
        )

        data = safe_json_loads(raw)
        questions = data.get("questions") if isinstance(data, dict) else None
        if not isinstance(questions, list):
            questions = []

        cleaned: List[str] = []
        for q in questions:
            q = clean_text(str(q))
            if not q:
                continue
            if not q.endswith("?"):
                q += "?"
            cleaned.append(q)

        return cleaned[:4]

    def analyze_image(
        self,
        image_path: str,
        *,
        user_question: str = "",
        language: str = "en",
    ) -> ImageInsight:
        client = self._require_client()

        max_side = env_int("IMAGE_MAX_SIDE", 1600)
        prepared_path = resize_image_for_vision(image_path, max_side=max_side)
        ocr_text = local_ocr_text(prepared_path)

        base_prompt = f"""
Analyze the image for a healthcare chatbot.

What to return:
- description of the scene or objects
- visible text if any
- medical relevance if any
- possible concerns
- concise notes for later reasoning

User question: {user_question or "No specific question"}
Language context: {language}

If the image is not medical, still describe it clearly.
If you are unsure, say so.
Return valid JSON with keys:
description, visible_text, medical_relevance, objects, notes
""".strip()

        if ocr_text:
            base_prompt += f"\n\nLocal OCR text already extracted:\n{ocr_text}"

        parsed: Dict[str, Any] = {}
        raw_text = ""

        try:
            data_url = image_to_data_url(prepared_path)
            for use_json in (True, False):
                try:
                    kwargs: Dict[str, Any] = {
                        "model": self.vision_model,
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": base_prompt},
                                    {"type": "image_url", "image_url": {"url": data_url}},
                                ],
                            }
                        ],
                        "temperature": 0.15,
                        "top_p": 1,
                        "max_completion_tokens": 450,
                    }
                    if use_json:
                        kwargs["response_format"] = {"type": "json_object"}

                    resp = client.chat.completions.create(**kwargs)
                    raw_text = resp.choices[0].message.content or ""
                    parsed = safe_json_loads(raw_text)
                    if parsed:
                        break
                except Exception:
                    continue
        finally:
            try:
                os.remove(prepared_path)
            except Exception:
                pass

        description = clean_text(str(parsed.get("description", "") if parsed else ""))
        visible_text = clean_text(str(parsed.get("visible_text", "") if parsed else ""))
        medical_relevance = clean_text(str(parsed.get("medical_relevance", "") if parsed else ""))
        notes = clean_text(str(parsed.get("notes", "") if parsed else ""))

        objects = parsed.get("objects", []) if isinstance(parsed, dict) else []
        if not isinstance(objects, list):
            objects = []

        if not description and raw_text and not parsed:
            description = clean_text(raw_text[:1200])

        if not description and not visible_text and not medical_relevance:
            description = "Image analysis completed with limited details available."
            if not visible_text:
                visible_text = ocr_text or ""
            if not notes:
                notes = "Fallback analysis was used because the vision response could not be fully parsed."

        if not visible_text and ocr_text:
            visible_text = ocr_text

        return ImageInsight(
            filename=Path(image_path).name,
            description=description,
            visible_text=visible_text,
            medical_relevance=medical_relevance,
            objects=[clean_text(str(x)) for x in objects if clean_text(str(x))],
            notes=notes,
            raw=parsed or {"raw": raw_text},
        )

    def transcribe_audio_file(self, audio_path: str, language: Optional[str] = None) -> str:
        client = self._require_client()

        with open(audio_path, "rb") as f:
            kwargs: Dict[str, Any] = {
                "file": f,
                "model": self.stt_model,
                "temperature": 0.0,
            }
            if language:
                kwargs["language"] = language

            transcription = client.audio.transcriptions.create(**kwargs)

        if hasattr(transcription, "text"):
            return clean_text(transcription.text)

        return clean_text(str(transcription))


class MemoryStore:
    def __init__(self, db_path: str, max_history_messages: int = 12) -> None:
        self.db_path = Path(db_path)
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self.max_history_messages = max_history_messages
        self._init_db()

    def _connect(self):
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn

    def _table_columns(self, conn, table: str) -> set[str]:
        rows = conn.execute(f"PRAGMA table_info({table})").fetchall()
        return {row["name"] for row in rows}

    def _ensure_turn_cache_columns(self, conn) -> None:
        cols = self._table_columns(conn, "turn_cache")
        if "assistant_answer" not in cols:
            conn.execute("ALTER TABLE turn_cache ADD COLUMN assistant_answer TEXT DEFAULT ''")

    def _init_db(self) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS sessions (
                    session_id TEXT PRIMARY KEY,
                    summary TEXT DEFAULT '',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS messages (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id TEXT NOT NULL,
                    role TEXT NOT NULL,
                    content TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            conn.execute(
                """
                CREATE INDEX IF NOT EXISTS idx_messages_session_id_id
                ON messages(session_id, id)
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS turn_cache (
                    session_id TEXT PRIMARY KEY,
                    user_message TEXT DEFAULT '',
                    assistant_answer TEXT DEFAULT '',
                    file_context TEXT DEFAULT '',
                    image_context TEXT DEFAULT '',
                    image_ocr TEXT DEFAULT '',
                    attachments_json TEXT DEFAULT '[]',
                    endpoint TEXT DEFAULT '',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            self._ensure_turn_cache_columns(conn)

    def ensure_session(self, session_id: str) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                INSERT OR IGNORE INTO sessions (session_id, summary)
                VALUES (?, '')
                """,
                (session_id,),
            )

    def get_summary(self, session_id: str) -> str:
        self.ensure_session(session_id)
        with self._connect() as conn:
            row = conn.execute(
                "SELECT summary FROM sessions WHERE session_id = ?",
                (session_id,),
            ).fetchone()
        return clean_text(row["summary"] if row else "")

    def set_summary(self, session_id: str, summary: str) -> None:
        self.ensure_session(session_id)
        summary = clean_text(summary)
        with self._connect() as conn:
            conn.execute(
                """
                UPDATE sessions
                SET summary = ?, updated_at = CURRENT_TIMESTAMP
                WHERE session_id = ?
                """,
                (summary, session_id),
            )

    def add_message(self, session_id: str, role: str, content: str) -> None:
        content = clean_text(content)
        if not content:
            return

        self.ensure_session(session_id)
        with self._connect() as conn:
            conn.execute(
                "INSERT INTO messages (session_id, role, content) VALUES (?, ?, ?)",
                (session_id, role, content),
            )

    def get_messages(self, session_id: str) -> List[Dict[str, Any]]:
        self.ensure_session(session_id)
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT id, role, content
                FROM messages
                WHERE session_id = ?
                ORDER BY id ASC
                """,
                (session_id,),
            ).fetchall()

        return [
            {"id": row["id"], "role": row["role"], "content": row["content"]}
            for row in rows
        ]

    def get_recent_messages(self, session_id: str, limit: int = 12) -> List[Dict[str, str]]:
        rows = self.get_messages(session_id)
        rows = rows[-limit:] if limit > 0 else rows
        return [{"role": row["role"], "content": row["content"]} for row in rows]

    def count_messages(self, session_id: str) -> int:
        self.ensure_session(session_id)
        with self._connect() as conn:
            row = conn.execute(
                "SELECT COUNT(*) AS count FROM messages WHERE session_id = ?",
                (session_id,),
            ).fetchone()
        return int(row["count"] if row else 0)

    def clear(self, session_id: str) -> None:
        self.ensure_session(session_id)
        with self._connect() as conn:
            conn.execute("DELETE FROM messages WHERE session_id = ?", (session_id,))
            conn.execute(
                "UPDATE sessions SET summary = '', updated_at = CURRENT_TIMESTAMP WHERE session_id = ?",
                (session_id,),
            )

    def get_context(self, session_id: str, keep_last: int = 12) -> Tuple[str, List[Dict[str, str]]]:
        summary = self.get_summary(session_id)
        recent = self.get_recent_messages(session_id, keep_last)
        return summary, recent

    def set_turn_cache(
        self,
        session_id: str,
        *,
        user_message: str = "",
        assistant_answer: str = "",
        file_context: str = "",
        image_context: str = "",
        image_ocr: str = "",
        attachments: Optional[List[Dict[str, Any]]] = None,
        endpoint: str = "",
    ) -> None:
        self.ensure_session(session_id)
        payload = json.dumps(attachments or [], ensure_ascii=False)

        with self._connect() as conn:
            self._ensure_turn_cache_columns(conn)
            conn.execute(
                """
                INSERT INTO turn_cache (
                    session_id, user_message, assistant_answer, file_context, image_context, image_ocr,
                    attachments_json, endpoint, created_at, updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                ON CONFLICT(session_id) DO UPDATE SET
                    user_message=excluded.user_message,
                    assistant_answer=excluded.assistant_answer,
                    file_context=excluded.file_context,
                    image_context=excluded.image_context,
                    image_ocr=excluded.image_ocr,
                    attachments_json=excluded.attachments_json,
                    endpoint=excluded.endpoint,
                    updated_at=CURRENT_TIMESTAMP
                """,
                (
                    session_id,
                    clean_text(user_message),
                    clean_text(assistant_answer),
                    clean_text(file_context),
                    clean_text(image_context),
                    clean_text(image_ocr),
                    payload,
                    endpoint,
                ),
            )

    def get_turn_cache(self, session_id: str) -> Dict[str, Any]:
        self.ensure_session(session_id)
        with self._connect() as conn:
            self._ensure_turn_cache_columns(conn)
            row = conn.execute(
                """
                SELECT user_message, assistant_answer, file_context, image_context, image_ocr,
                       attachments_json, endpoint
                FROM turn_cache
                WHERE session_id = ?
                """,
                (session_id,),
            ).fetchone()

        if not row:
            return {}

        try:
            attachments = json.loads(row["attachments_json"] or "[]")
        except Exception:
            attachments = []

        return {
            "user_message": clean_text(row["user_message"] or ""),
            "assistant_answer": clean_text(row["assistant_answer"] or ""),
            "file_context": clean_text(row["file_context"] or ""),
            "image_context": clean_text(row["image_context"] or ""),
            "image_ocr": clean_text(row["image_ocr"] or ""),
            "attachments": attachments if isinstance(attachments, list) else [],
            "endpoint": clean_text(row["endpoint"] or ""),
        }

    def clear_turn_cache(self, session_id: str) -> None:
        self.ensure_session(session_id)
        with self._connect() as conn:
            conn.execute("DELETE FROM turn_cache WHERE session_id = ?", (session_id,))

    def replace_last_assistant(self, session_id: str, assistant_text: str) -> None:
        assistant_text = clean_text(assistant_text)
        if not assistant_text:
            return

        self.ensure_session(session_id)
        with self._connect() as conn:
            row = conn.execute(
                """
                SELECT id
                FROM messages
                WHERE session_id = ? AND role = 'assistant'
                ORDER BY id DESC
                LIMIT 1
                """,
                (session_id,),
            ).fetchone()

            if row:
                conn.execute("DELETE FROM messages WHERE id = ?", (row["id"],))

            conn.execute(
                "INSERT INTO messages (session_id, role, content) VALUES (?, 'assistant', ?)",
                (session_id, assistant_text),
            )

    def compact_if_needed(
        self,
        session_id: str,
        groq_service: GroqService,
        *,
        language: str = "en",
        keep_last: int = 12,
        trigger_after: int = 28,
    ) -> bool:
        total = self.count_messages(session_id)
        if total <= trigger_after:
            return False

        rows = self.get_messages(session_id)
        if len(rows) <= keep_last:
            return False

        older = rows[:-keep_last]
        if not older:
            return False

        older_blob = "\n".join(f"{row['role'].upper()}: {row['content']}" for row in older)
        current_summary = self.get_summary(session_id)

        try:
            new_summary = groq_service.compress_chat_memory(
                previous_summary=current_summary,
                older_messages=older_blob,
                language=language,
            )
        except Exception:
            return False

        if not new_summary:
            return False

        ids_to_delete = [row["id"] for row in older]
        if not ids_to_delete:
            return False

        with self._connect() as conn:
            placeholders = ",".join(["?"] * len(ids_to_delete))
            conn.execute(f"DELETE FROM messages WHERE id IN ({placeholders})", ids_to_delete)
            conn.execute(
                """
                UPDATE sessions
                SET summary = ?, updated_at = CURRENT_TIMESTAMP
                WHERE session_id = ?
                """,
                (new_summary, session_id),
            )

        return True


def safe_delete(path: Optional[str]) -> None:
    if not path:
        return
    try:
        p = Path(path)
        if p.exists() and p.is_file():
            p.unlink()
    except Exception:
        pass