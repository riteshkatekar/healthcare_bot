
from __future__ import annotations

import csv
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from PIL import Image, ImageFilter, ImageOps
from pypdf import PdfReader

from .language_service import (
    clean_text,
    chunk_text,
    detect_language,
    env_int,
    env_str,
    safe_snippet,
)

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import cv2  # type: ignore
except Exception:
    cv2 = None

try:
    import pytesseract
except Exception:
    pytesseract = None


ALLOWED_TEXT_EXTS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".html", ".htm", ".xml", ".log"}
ALLOWED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}
ALLOWED_AUDIO_EXTS = {".wav", ".mp3", ".m4a", ".mp4", ".mpeg", ".mpga", ".ogg", ".webm", ".flac"}

WHITESPACE_RE = re.compile(r"\s+")
HTML_TAG_RE = re.compile(r"<[^>]+>")
JSON_BLOCK_RE = re.compile(r"\{.*\}", re.DOTALL)

# Tesseract OCR Configuration
TESSERACT_CMD = os.getenv(
    "TESSERACT_CMD",
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
).strip()

TESSERACT_DEFAULT_LANGS = os.getenv(
    "TESSERACT_LANGS",
    "eng+hin+mar",
).strip() or "eng+hin+mar"

TESSERACT_CONFIG = os.getenv(
    "TESSERACT_CONFIG",
    "--oem 3 --psm 6",
).strip()

if pytesseract is not None and TESSERACT_CMD:
    try:
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
    except Exception:
        pass


def _extract_pdf_text_pages(file_path: str) -> List[str]:
    pages: List[str] = []
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
    except Exception:
        return []
    return pages


def _pdf_text_stats(pages: Sequence[str]) -> Tuple[int, int]:
    nonempty = 0
    total_chars = 0
    for page in pages:
        txt = clean_text(page)
        if txt:
            nonempty += 1
            total_chars += len(txt)
    return nonempty, total_chars


def detect_pdf_type(file_path: str, *, min_text_chars: int = 40) -> str:
    pages = _extract_pdf_text_pages(file_path)
    if not pages:
        return "scanned"

    nonempty, total_chars = _pdf_text_stats(pages)
    if total_chars >= min_text_chars and nonempty >= max(1, len(pages) // 2):
        return "text"

    if nonempty > 0:
        return "mixed"

    return "scanned"


def _sanitize_pdf_page_for_ocr(page_image: Image.Image) -> Image.Image:
    img = ImageOps.exif_transpose(page_image).convert("L")
    img = ImageOps.autocontrast(img)

    if cv2 is not None:
        try:
            import numpy as np  # type: ignore
            arr = np.array(img)
            arr = cv2.GaussianBlur(arr, (3, 3), 0)
            arr = cv2.adaptiveThreshold(
                arr,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                31,
                11,
            )
            return Image.fromarray(arr)
        except Exception:
            pass

    img = img.filter(ImageFilter.SHARPEN)
    return img


def resize_image_for_vision(input_path: str, max_side: int = 1600) -> str:
    img = Image.open(input_path)
    img = ImageOps.exif_transpose(img).convert("RGB")

    w, h = img.size
    longest = max(w, h)
    if longest > max_side:
        scale = max_side / float(longest)
        new_size = (max(1, int(w * scale)), max(1, int(h * scale)))
        img = img.resize(new_size, Image.LANCZOS)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    tmp_path = tmp.name
    tmp.close()
    img.save(tmp_path, format="JPEG", quality=92, optimize=True)
    return tmp_path


def local_ocr_text(image_path: str, languages: Optional[str] = None) -> str:
    if pytesseract is None:
        return ""
    try:
        img = Image.open(image_path)
        img = _sanitize_pdf_page_for_ocr(img)
        lang = (languages or TESSERACT_DEFAULT_LANGS).strip() or "eng"
        text = pytesseract.image_to_string(img, lang=lang, config=TESSERACT_CONFIG)
        return clean_text(text)
    except Exception:
        try:
            img = Image.open(image_path)
            img = _sanitize_pdf_page_for_ocr(img)
            text = pytesseract.image_to_string(img, config=TESSERACT_CONFIG)
            return clean_text(text)
        except Exception:
            return ""


def render_pdf_pages(file_path: str, *, max_pages: Optional[int] = None, zoom: float = 2.0) -> List[str]:
    if fitz is None:
        raise RuntimeError("PyMuPDF is required to render PDF pages. Install pymupdf.")

    doc = fitz.open(file_path)
    paths: List[str] = []
    try:
        page_count = len(doc)
        limit = page_count if max_pages is None else min(page_count, max_pages)
        matrix = fitz.Matrix(zoom, zoom)

        for index in range(limit):
            page = doc.load_page(index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f"_page_{index + 1}.png")
            tmp_path = tmp.name
            tmp.close()
            pix.save(tmp_path)
            paths.append(tmp_path)
    finally:
        doc.close()

    return paths


def ocr_pdf_pages(page_paths: Sequence[str], *, languages: Optional[str] = None) -> str:
    pieces: List[str] = []
    lang = (languages or TESSERACT_DEFAULT_LANGS).strip() or "eng"

    for idx, page_path in enumerate(page_paths, start=1):
        text = local_ocr_text(page_path, languages=lang)
        if text:
            pieces.append(f"[Page {idx}]\n{text}")
        else:
            pieces.append(f"[Page {idx}]\n")

    return clean_text("\n".join(pieces))


def _extract_medical_fields(text: str) -> Dict[str, Any]:
    source = clean_text(text)
    lines = [clean_text(line) for line in source.splitlines()]
    lines = [line for line in lines if line]

    def _find(patterns: Sequence[str]) -> str:
        for pattern in patterns:
            m = re.search(pattern, source, flags=re.IGNORECASE | re.MULTILINE)
            if m:
                value = m.group(1) if m.groups() else m.group(0)
                return clean_text(value)
        return ""

    patient_name = _find([
        r"(?:patient(?: name)?|name)\s*[:\-]\s*([A-Za-z][A-Za-z .'-]{2,})",
        r"(?:patient(?: name)?|name)\s*[:\-]\s*([^\n\r]+)",
    ])
    if not patient_name and lines:
        for line in lines[:8]:
            if any(k in line.lower() for k in ("patient", "name")) and len(line) < 120:
                patient_name = clean_text(line.split(":")[-1])
                break

    date = _find([
        r"(?:date|dated)\s*[:\-]\s*([0-9]{1,2}[\/\-.][0-9]{1,2}[\/\-.][0-9]{2,4})",
        r"(?:date|dated)\s*[:\-]\s*([A-Za-z]{3,9}\s+[0-9]{1,2},\s*[0-9]{4})",
        r"([0-9]{4}[\/\-.][0-9]{1,2}[\/\-.][0-9]{1,2})",
    ])

    age = _find([
        r"(?:age)\s*[:\-]\s*([0-9]{1,3})",
        r"([0-9]{1,3})\s*(?:years?|yrs?)\b",
    ])

    sex = _find([
        r"(?:sex|gender)\s*[:\-]\s*(male|female|m|f|transgender|other)\b",
    ])
    if sex:
        sex = sex.capitalize()

    doctor = _find([
        r"(?:doctor|dr\.?|physician)\s*[:\-]\s*([A-Za-z][A-Za-z .'-]{2,})",
    ])

    diagnosis = _find([
        r"(?:diagnosis|impression|assessment|provisional diagnosis)\s*[:\-]\s*([^\n\r]+)",
    ])

    symptoms: List[str] = []
    symptom_patterns = [
        r"(?:complains of|c/o|presenting with|symptoms?[:\-])\s*([^\n\r]+)",
        r"(?:chief complaint|cc)[:\-]\s*([^\n\r]+)",
    ]
    for pat in symptom_patterns:
        m = re.search(pat, source, flags=re.IGNORECASE | re.MULTILINE)
        if m:
            chunk = clean_text(m.group(1))
            symptoms.extend([clean_text(x) for x in re.split(r"[;,/]| and ", chunk) if clean_text(x)])

    medicines: List[Dict[str, str]] = []
    med_line_patterns = [
        r"(?:medicines?|drugs?|prescribed|rx)\s*[:\-]\s*([^\n\r]+)",
        r"(?:tab(?:let)?|cap(?:sule)?|syrup|inj(?:ection)?)\s+([^\n\r]+)",
    ]
    for pat in med_line_patterns:
        for m in re.finditer(pat, source, flags=re.IGNORECASE | re.MULTILINE):
            block = clean_text(m.group(1))
            if not block:
                continue
            for part in re.split(r"[;,]| and ", block):
                part = clean_text(part)
                if not part:
                    continue
                medicines.append({"name": part, "dose": "", "frequency": ""})

    lab_values: List[Dict[str, str]] = []
    lab_patterns = {
        "bp": r"\b(?:bp|blood pressure)\b[:\-\s]*([0-9]{2,3}\s*/\s*[0-9]{2,3})",
        "sugar": r"\b(?:sugar|glucose|rbs|fbs)\b[:\-\s]*([0-9]+(?:\.[0-9]+)?)\s*(mg\/dl|mmol\/l)?",
        "hb": r"\b(?:hb|hgb|hemoglobin)\b[:\-\s]*([0-9]+(?:\.[0-9]+)?)\s*(g\/dl)?",
        "spo2": r"\b(?:spo2|oxygen saturation)\b[:\-\s]*([0-9]+(?:\.[0-9]+)?)\s*%?",
        "pulse": r"\b(?:pulse|pr)\b[:\-\s]*([0-9]+)\s*(bpm)?",
        "temperature": r"\b(?:temp|temperature|fever)\b[:\-\s]*([0-9]+(?:\.[0-9]+)?)\s*(?:c|°c|f|°f)?",
    }
    for name, pat in lab_patterns.items():
        for m in re.finditer(pat, source, flags=re.IGNORECASE | re.MULTILINE):
            value = clean_text(m.group(1))
            unit = clean_text(m.group(2)) if m.lastindex and m.lastindex >= 2 else ""
            lab_values.append({"test": name.upper(), "value": value, "unit": unit})

    return {
        "patient_name": patient_name,
        "date": date,
        "age": age,
        "sex": sex,
        "doctor": doctor,
        "diagnosis": diagnosis,
        "symptoms": symptoms,
        "medicines": medicines[:20],
        "lab_values": lab_values[:25],
        "raw_text": source,
    }


def structure_medical_values(raw_text: str, *, language: str = "en") -> Dict[str, Any]:
    text = clean_text(raw_text)
    fields = _extract_medical_fields(text)

    prompt_facts = []
    if fields.get("patient_name"):
        prompt_facts.append(f"Patient name: {fields['patient_name']}")
    if fields.get("date"):
        prompt_facts.append(f"Date: {fields['date']}")
    if fields.get("age"):
        prompt_facts.append(f"Age: {fields['age']}")
    if fields.get("sex"):
        prompt_facts.append(f"Sex: {fields['sex']}")
    if fields.get("doctor"):
        prompt_facts.append(f"Doctor: {fields['doctor']}")
    if fields.get("diagnosis"):
        prompt_facts.append(f"Diagnosis: {fields['diagnosis']}")
    if fields.get("symptoms"):
        prompt_facts.append("Symptoms: " + ", ".join(fields["symptoms"]))
    if fields.get("medicines"):
        med_text = "; ".join(
            f"{m.get('name','')}{(' ' + m.get('dose','')).strip() if m.get('dose') else ''}".strip()
            for m in fields["medicines"]
        )
        prompt_facts.append("Medicines: " + med_text)
    if fields.get("lab_values"):
        lv_text = "; ".join(
            f"{lv.get('test','')}: {lv.get('value','')}{(' ' + lv.get('unit','')).strip() if lv.get('unit') else ''}"
            for lv in fields["lab_values"]
        )
        prompt_facts.append("Lab values: " + lv_text)

    structured_summary = "\n".join(prompt_facts).strip()
    fields["structured_summary"] = structured_summary
    fields["language_hint"] = language
    return fields


def _pdf_page_context(page_no: int, page_text: str, structured: Optional[Dict[str, Any]] = None) -> str:
    blocks = [f"[Page {page_no}]"]
    page_text = clean_text(page_text)
    if page_text:
        blocks.append(page_text)
    if structured and structured.get("structured_summary"):
        blocks.append("Structured medical facts:\n" + structured["structured_summary"])
    return "\n".join(blocks).strip()


def prepare_pdf_context(file_path: str, filename: str, user_message: str = "") -> Tuple[str, Dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    if suffix != ".pdf":
        raise ValueError("prepare_pdf_context() expects a PDF file.")

    pages = _extract_pdf_text_pages(file_path)
    pdf_type = detect_pdf_type(file_path)
    language_hint = detect_language(user_message or "")

    attachment: Dict[str, Any] = {
        "type": "file",
        "filename": filename,
        "status": "uploaded",
        "file_type": "pdf",
        "pdf_type": pdf_type,
        "path": Path(file_path).name,
    }

    text_layer = clean_text("\n".join(pages))
    if pdf_type == "text" and text_layer:
        structured = structure_medical_values(text_layer, language=language_hint)
        condensed = text_layer
        if structured.get("structured_summary"):
            condensed = structured["structured_summary"] + "\n\n" + safe_snippet(text_layer, 9000)
        attachment["status"] = "processed"
        attachment["summary"] = safe_snippet(condensed, 600)
        attachment["structured"] = structured
        return condensed, attachment

    max_pdf_pages = env_int("PDF_OCR_MAX_PAGES", 30)
    max_vision_pages = env_int("PDF_VISION_MAX_PAGES", 12)
    zoom = float(env_str("PDF_RENDER_ZOOM", "2.0"))

    page_paths: List[str] = []
    ocr_text = ""
    try:
        page_paths = render_pdf_pages(
            file_path,
            max_pages=max_vision_pages if pdf_type == "mixed" else max_pdf_pages,
            zoom=zoom,
        )
        ocr_text = ocr_pdf_pages(page_paths, languages=TESSERACT_DEFAULT_LANGS)
    finally:
        for p in page_paths:
            try:
                Path(p).unlink(missing_ok=True)
            except Exception:
                pass

    merged = "\n\n".join(part for part in [text_layer, ocr_text] if clean_text(part)).strip()
    structured = structure_medical_values(merged, language=language_hint)

    attachment["status"] = "processed"
    attachment["summary"] = safe_snippet(structured.get("structured_summary") or merged, 600)
    attachment["structured"] = structured

    if structured.get("structured_summary"):
        output = structured["structured_summary"] + "\n\n" + safe_snippet(merged, 9000)
    else:
        output = merged

    return clean_text(output), attachment


def extract_text_from_file(file_path: str, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    try:
        if suffix == ".pdf":
            text, _attachment = prepare_pdf_context(file_path, filename)
            return text

        if suffix == ".docx":
            doc = Document(file_path)
            parts: List[str] = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return clean_text("\n".join(parts))

        if suffix == ".txt":
            return clean_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"))

        if suffix == ".md":
            return clean_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"))

        if suffix == ".csv":
            rows: List[str] = []
            with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    row = [cell.strip() for cell in row if cell is not None and str(cell).strip()]
                    if row:
                        rows.append(" | ".join(row))
            return clean_text("\n".join(rows))

        if suffix == ".json":
            raw = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            try:
                parsed = json.loads(raw)
                return clean_text(json.dumps(parsed, indent=2, ensure_ascii=False))
            except Exception:
                return clean_text(raw)

        if suffix in {".html", ".htm", ".xml", ".log"}:
            raw = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            if suffix in {".html", ".htm"}:
                raw = HTML_TAG_RE.sub(" ", raw)
            return clean_text(raw)

    except Exception:
        return ""

    return ""
